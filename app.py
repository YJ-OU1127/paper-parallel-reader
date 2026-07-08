import html
import json
import math
import re
from datetime import datetime
from io import BytesIO

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Pt


APP_NAME = "Paper Parallel Reader"
APP_VERSION = "1.5.9-translation-typography-match"

STATUS_OPTIONS = ["未確認", "確認済み", "要修正", "修正済み"]
GLOSSARY_CATEGORIES = ["専門用語", "理論概念", "方法", "固有名詞", "その他"]
WORD_EXPORT_OPTIONS = [
    "2列表（元文｜訳文）",
    "縦並び（元文→訳文）",
    "訳文のみ",
    "確認用（ID・状態・メモつき）",
]

DEFAULT_VIEWER_HEIGHT = 560
READ_EDITOR_HEIGHT = 620

st.set_page_config(page_title=APP_NAME, layout="wide")


# ============================================================
# デザイン
# ============================================================
def inject_style():
    st.markdown(
        """
<style>
:root {
  --main-blue: #5A6F97;
  --deep-blue: #25243D;
  --soft-blue: #EEF2F7;
  --main-green: #2E8B7D;
  --soft-green: #EEF7F3;
  --paper: #FCF7EE;
  --ink: #26272F;
  --muted: #756F68;
  --line: #E8DCCB;
  --rose: #C46A6A;
  --apricot: #E7A76D;
  --sand: #F3E6C8;
  --lavender: #C8BEDD;
  --ppr-unchecked: #D8CFC1;
  --ppr-needsfix: #D67A59;
  --ppr-confirmed: #6078A8;
  --ppr-revised: #2E8B7D;
}

.block-container {
  padding-top: 1.2rem;
  padding-bottom: 3rem;
}

.stApp {
  background:
    radial-gradient(circle at 74% 0%, rgba(200,190,221,.23) 0, rgba(200,190,221,0) 28%),
    linear-gradient(180deg, #FFFCF7 0%, #F9F5EC 100%);
}

.ppr-hero {
  background:
    radial-gradient(circle at 92% 10%, rgba(231, 167, 109, .38) 0, rgba(231, 167, 109, 0) 30%),
    linear-gradient(135deg, #25243D 0%, #4E536C 52%, #B07368 100%);
  color: #FFF9EF;
  border-radius: 24px;
  padding: 24px 28px;
  box-shadow: 0 18px 42px rgba(45, 36, 36, .20);
  margin-bottom: 18px;
}
.ppr-hero h1 {
  margin: 0 0 6px 0;
  font-size: 2.05rem;
  letter-spacing: .01em;
}
.ppr-hero p {
  margin: 0;
  color: rgba(255,249,239,.88);
  font-size: .98rem;
}
.ppr-badge {
  display:inline-block;
  padding: 4px 10px;
  border-radius: 999px;
  background: rgba(255,249,239,.14);
  border: 1px solid rgba(255,249,239,.28);
  margin-right: 6px;
  margin-top: 10px;
  font-size: .82rem;
}
.ppr-card {
  background: rgba(255,253,247,.92);
  border: 1px solid var(--line);
  border-radius: 20px;
  padding: 14px 16px;
  box-shadow: 0 12px 28px rgba(45,36,36,.06);
}
.ppr-subtle {
  color: var(--muted);
  font-size: .92rem;
}
.ppr-chip {
  display:inline-block;
  padding: 4px 10px;
  border-radius: 999px;
  border: 1px solid var(--line);
  background: var(--paper);
  margin: 2px 4px 2px 0;
  font-size: .82rem;
}
.ppr-chip-blue { background: #EEF2F7; border-color:#D6DDEA; color:#4E6388; }
.ppr-chip-green { background: #EEF7F3; border-color:#C8E1D7; color:#2E8B7D; }
.ppr-chip-yellow { background:#FFF1D7; border-color:#EBC78E; color:#95603A; }
.ppr-chip-red { background:#F8E8E2; border-color:#E6B5A4; color:#A14F44; }

[data-testid="stMetricValue"] { color: var(--deep-blue); }

.stTabs [data-baseweb="tab-list"] {
  gap: 8px;
  border-bottom: 1px solid var(--line);
}
.stTabs [data-baseweb="tab"] {
  border-radius: 12px 12px 0 0;
  padding: 10px 14px;
  font-weight: 700;
}
.stTabs [aria-selected="true"] {
  background: var(--soft-blue);
  color: var(--main-blue);
}

div.stButton > button:first-child,
div.stDownloadButton > button:first-child {
  border-radius: 13px;
  border: 1px solid #D9C8B8;
  font-weight: 750;
  background: rgba(255,253,247,.72);
  color: var(--ink);
}
div.stButton > button:first-child:hover,
div.stDownloadButton > button:first-child:hover {
  border-color: #C46A6A;
  color: #8C4B4B;
  background: #FFF9F4;
}

section[data-testid="stSidebar"] {
  background:
    radial-gradient(circle at 20% 4%, rgba(231,167,109,.18) 0, rgba(231,167,109,0) 34%),
    linear-gradient(180deg, #FBF7EF 0%, #F3EFE6 100%);
}

.ppr-progress-card {
  background: rgba(255,255,255,.88);
  border: 1px solid var(--line);
  border-radius: 18px;
  padding: 14px 14px 16px;
  box-shadow: 0 10px 24px rgba(31,41,51,.06);
}
.ppr-progress-head {
  display: flex;
  justify-content: space-between;
  align-items: baseline;
  gap: 8px;
  margin-bottom: 12px;
}
.ppr-progress-title {
  font-weight: 900;
  color: var(--deep-blue);
  letter-spacing: .02em;
}
.ppr-progress-percent {
  font-weight: 900;
  color: var(--ppr-revised);
  font-size: 1.12rem;
  font-variant-numeric: tabular-nums;
}
.ppr-progress-body {
  display: flex;
  justify-content: center;
  align-items: center;
}
.ppr-progress-stack {
  width: 112px;
  height: 104px;
  position: relative;
  display: flex;
  flex-direction: column;
  border-radius: 18px;
  background: #EDE5D8;
  box-shadow: inset 0 0 0 1px rgba(255,255,255,.58), 0 8px 18px rgba(31,41,51,.10);
  overflow: visible;
}
.ppr-progress-segment {
  position: relative;
  width: 100%;
  min-height: 4px;
  cursor: help;
  transition: filter .12s ease, transform .12s ease;
}
.ppr-progress-segment:hover {
  filter: brightness(.94) saturate(1.08);
  z-index: 15;
}
.ppr-progress-segment:first-child {
  border-radius: 18px 18px 0 0;
}
.ppr-progress-segment:last-child {
  border-radius: 0 0 18px 18px;
}
.ppr-progress-segment:only-child {
  border-radius: 18px;
}
.ppr-progress-unchecked { background: var(--ppr-unchecked); }
.ppr-progress-confirmed { background: var(--ppr-confirmed); }
.ppr-progress-needsfix { background: var(--ppr-needsfix); }
.ppr-progress-revised { background: var(--ppr-revised); }
.ppr-progress-tip {
  display: none;
  position: absolute;
  left: 50%;
  top: 50%;
  transform: translate(-50%, -50%);
  width: max-content;
  min-width: 154px;
  max-width: 190px;
  padding: 9px 11px;
  border-radius: 13px;
  background: rgba(24,34,45,.96);
  color: #fffdf7;
  font-size: 12.5px;
  line-height: 1.45;
  z-index: 50;
  box-shadow: 0 12px 26px rgba(31,41,51,.24);
  pointer-events: none;
  white-space: normal;
}
.ppr-progress-segment:hover .ppr-progress-tip {
  display: block;
}
.ppr-progress-tip-label {
  display: block;
  font-weight: 900;
  color: #f2eadb;
  margin-bottom: 2px;
}
.ppr-progress-tip-detail {
  display: block;
  color: rgba(255,253,247,.86);
  font-variant-numeric: tabular-nums;
}


/* 読む・修正：Japanese Translation の編集欄を Original English に近づける */
textarea[aria-label="Japanese Translation"] {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif !important;
  font-size: 15.5px !important;
  line-height: 1.86 !important;
  letter-spacing: .01em !important;
  color: #26272F !important;
  background: linear-gradient(180deg,#FFFDF7 0%,#FBF7EF 100%) !important;
  border: 1px solid #E8DCCB !important;
  border-radius: 18px !important;
  padding: 16px 18px !important;
  box-shadow: 0 10px 24px rgba(45,36,36,.05) !important;
}
textarea[aria-label="Japanese Translation"]:focus {
  border-color: #9A8FBD !important;
  box-shadow: 0 0 0 3px rgba(154,143,189,.18), 0 10px 24px rgba(45,36,36,.05) !important;
}
label:has(+ div textarea[aria-label="Japanese Translation"]) {
  color: #25243D !important;
  font-size: 16px !important;
  font-weight: 800 !important;
}

</style>
""",
        unsafe_allow_html=True,
    )


# ============================================================
# 基本ユーティリティ
# ============================================================
def init_state():
    defaults = {
        "rows": [],
        "glossary": [],
        "project_title": "paper_parallel_project",
        "unit": "段落",
        "selected_id": None,
        "active_page": "一覧・選択",
        "page_size": 20,
        "context_radius": 2,
        "show_settings": True,
        "show_json_loader": True,
        "show_text_adder": True,
        "last_filtered_ids": [],
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def now_string():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def safe_filename(text, fallback="paper_parallel"):
    text = str(text or fallback).strip() or fallback
    text = re.sub(r"[\\/:*?\"<>|\s]+", "_", text)
    return text[:80]


def truncate_text(text, max_chars=120):
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    return text if len(text) <= max_chars else text[:max_chars] + "…"


def estimate_english_words(text):
    return len(re.findall(r"[A-Za-z]+(?:[-'][A-Za-z]+)?", str(text or "")))


def get_final_translation(row):
    edited = str(row.get("修正訳", "")).strip()
    return edited if edited else str(row.get("原訳", "")).strip()


def normalize_status(status):
    status = str(status or "未確認")
    if status in STATUS_OPTIONS:
        return status
    if status == "手動":
        return "修正済み"
    return "未確認"


def clean_phrase(text):
    return re.sub(r"\s+", " ", str(text or "")).strip()


def lower_key(text):
    return clean_phrase(text).lower()


def next_item_id(items):
    ids = []
    for item in items:
        try:
            ids.append(int(item.get("id", 0)))
        except Exception:
            pass
    return max(ids or [0]) + 1


# ============================================================
# 分割関数
# ============================================================
def split_paragraphs(text):
    text = str(text or "").strip()
    if not text:
        return []
    paragraphs = re.split(r"\n\s*\n+", text)
    return [re.sub(r"\s+", " ", p.strip()) for p in paragraphs if p.strip()]


def split_english_sentences(text):
    text = re.sub(r"\s+", " ", str(text or "").strip())
    if not text:
        return []
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9\"'(\[])", text)
    return [s.strip() for s in sentences if s.strip()]


def split_japanese_sentences(text):
    text = re.sub(r"\s+", " ", str(text or "").strip())
    if not text:
        return []
    sentences = re.split(r"(?<=[。！？!?])\s*", text)
    return [s.strip() for s in sentences if s.strip()]


def segment_text(text, language, unit):
    if unit == "段落":
        return split_paragraphs(text)
    return split_english_sentences(text) if language == "英語" else split_japanese_sentences(text)


# ============================================================
# 行データ
# ============================================================
def make_rows_from_segments(english_segments, japanese_segments, section_title, start_id=1):
    section_title = str(section_title or "Section").strip() or "Section"
    row_count = max(len(english_segments), len(japanese_segments))
    rows = []
    for i in range(row_count):
        en = english_segments[i] if i < len(english_segments) else ""
        ja = japanese_segments[i] if i < len(japanese_segments) else ""
        rows.append(
            {
                "ID": start_id + i,
                "章": section_title,
                "章内ID": i + 1,
                "英文": en,
                "原訳": ja,
                "修正訳": ja,
                "状態": "未確認",
                "メモ": "",
                "更新日時": "",
            }
        )
    return rows


def normalize_row(row, fallback_id, fallback_section="Imported"):
    if "ID" in row and "英文" in row:
        original = str(row.get("原訳", row.get("和訳", row.get("訳文", ""))))
        edited = str(row.get("修正訳", row.get("和訳（編集用）", row.get("訳文", original))))
        return {
            "ID": int(row.get("ID") or fallback_id),
            "章": str(row.get("章", fallback_section) or fallback_section),
            "章内ID": int(row.get("章内ID") or fallback_id),
            "英文": str(row.get("英文", "")),
            "原訳": original,
            "修正訳": edited,
            "状態": normalize_status(row.get("状態", row.get("修正状態", "未確認"))),
            "メモ": str(row.get("メモ", "")),
            "更新日時": str(row.get("更新日時", "")),
        }

    original = str(row.get("和訳", row.get("訳文", "")))
    edited = str(row.get("和訳（編集用）", original))
    status = "修正済み" if row.get("修正状態") == "手動" else "未確認"
    return {
        "ID": int(row.get("英語ID") or fallback_id),
        "章": fallback_section,
        "章内ID": int(row.get("英語ID") or fallback_id),
        "英文": str(row.get("英文", "")),
        "原訳": original,
        "修正訳": edited,
        "状態": status,
        "メモ": "",
        "更新日時": "",
    }


def get_row_index_by_id(rows, row_id):
    for index, row in enumerate(rows):
        try:
            if int(row.get("ID", -1)) == int(row_id):
                return index
        except Exception:
            continue
    return None


def get_row_by_id(rows, row_id):
    index = get_row_index_by_id(rows, row_id)
    return rows[index] if index is not None else None


def renumber_rows(rows):
    section_counts = {}
    for index, row in enumerate(rows, start=1):
        section = str(row.get("章", "Section") or "Section")
        section_counts[section] = section_counts.get(section, 0) + 1
        row["ID"] = index
        row["章内ID"] = section_counts[section]
    return rows


# ============================================================
# 用語辞典
# ============================================================
def glossary_unique_key(item):
    return (
        lower_key(item.get("english", "")),
        clean_phrase(item.get("japanese", "")),
        str(item.get("category", "")),
    )


def cleanup_glossary():
    glossary = []
    seen = set()
    for item in st.session_state.get("glossary", []):
        english = clean_phrase(item.get("english", item.get("英語", "")))
        japanese = clean_phrase(item.get("japanese", item.get("日本語", "")))
        if not english and not japanese:
            continue
        fixed = {
            "id": int(item.get("id", len(glossary) + 1) or len(glossary) + 1),
            "english": english,
            "japanese": japanese,
            "category": str(item.get("category", item.get("カテゴリ", "専門用語")) or "専門用語"),
            "note": str(item.get("note", item.get("メモ", ""))),
            "created_at": str(item.get("created_at", "")),
        }
        key = glossary_unique_key(fixed)
        if key in seen:
            continue
        seen.add(key)
        glossary.append(fixed)

    for i, item in enumerate(glossary, start=1):
        item["id"] = i
    st.session_state["glossary"] = glossary


def add_unique_glossary(english, japanese, category, note):
    item = {
        "id": next_item_id(st.session_state["glossary"]),
        "english": clean_phrase(english),
        "japanese": clean_phrase(japanese),
        "category": category,
        "note": str(note or "").strip(),
        "created_at": now_string(),
    }
    if not item["english"] and not item["japanese"]:
        return False, "英語用語または日本語訳を入力してください。"

    existing_keys = {glossary_unique_key(g) for g in st.session_state["glossary"]}
    if glossary_unique_key(item) in existing_keys:
        return False, "同じ用語はすでに登録されています。"

    st.session_state["glossary"].append(item)
    cleanup_glossary()
    return True, "用語辞典に追加しました。"


def has_glossary_hit(row):
    en = str(row.get("英文", ""))
    for item in st.session_state.get("glossary", []):
        english = clean_phrase(item.get("english", ""))
        if english and re.search(make_term_pattern(english), en, flags=re.IGNORECASE):
            return True
    return False


def make_term_pattern(term):
    escaped = re.escape(clean_phrase(term))
    if not escaped:
        return r"a^"
    left = r"(?<![A-Za-z0-9_-])" if re.match(r"^[A-Za-z0-9]", term) else ""
    right = r"(?![A-Za-z0-9_-])" if re.search(r"[A-Za-z0-9]$", term) else ""
    return f"{left}{escaped}{right}"


# ============================================================
# JSON読み込み・保存
# ============================================================
def load_rows_from_json(uploaded_file):
    data = json.loads(uploaded_file.getvalue().decode("utf-8"))
    project_title = data.get("project_title") or data.get("title") or "paper_parallel_project"
    unit = data.get("unit", "段落")

    if "rows" in data:
        rows = [normalize_row(row, i + 1) for i, row in enumerate(data.get("rows", []))]
    elif "alignment_results" in data:
        rows = [normalize_row(row, i + 1, "Imported") for i, row in enumerate(data.get("alignment_results", []))]
    else:
        raise ValueError("rows または alignment_results が見つかりません。")

    glossary = []
    for i, item in enumerate(data.get("glossary", []), start=1):
        glossary.append(
            {
                "id": int(item.get("id", i) or i),
                "english": clean_phrase(item.get("english", item.get("英語", ""))),
                "japanese": clean_phrase(item.get("japanese", item.get("日本語", ""))),
                "category": str(item.get("category", item.get("カテゴリ", "専門用語")) or "専門用語"),
                "note": str(item.get("note", item.get("メモ", ""))),
                "created_at": str(item.get("created_at", "")),
            }
        )

    st.session_state["rows"] = renumber_rows(rows)
    st.session_state["glossary"] = glossary
    cleanup_glossary()
    return project_title, unit, st.session_state["rows"]


def make_json_data(project_title, unit, rows):
    cleanup_glossary()
    data = {
        "app_name": APP_NAME,
        "version": APP_VERSION,
        "saved_at": now_string(),
        "project_title": project_title,
        "unit": unit,
        "rows": rows,
        "glossary": st.session_state.get("glossary", []),
    }
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


# ============================================================
# 表示用データ
# ============================================================
def progress_counts(rows):
    total = len(rows)
    counts = {status: 0 for status in STATUS_OPTIONS}
    for row in rows:
        counts[normalize_status(row.get("状態"))] += 1
    done = counts["確認済み"] + counts["修正済み"]
    percent = round((done / total) * 100, 1) if total else 0.0
    return {
        "total": total,
        "unchecked": counts["未確認"],
        "confirmed": counts["確認済み"],
        "needsfix": counts["要修正"],
        "revised": counts["修正済み"],
        "done": done,
        "percent": percent,
    }


def render_sidebar_progress_card(rows):
    # サイドバー内に、軽量な縦向き進捗グラフを表示する。
    # Streamlit の markdown HTML では div が崩れて閉じタグが文字として見える場合があるため、
    # components.html の小さな iframe 内で完結させる。
    counts = progress_counts(rows)
    total = counts["total"]
    bar_height_px = 318
    min_visible_px = 8

    def pct(value):
        return (value / total * 100) if total else 0.0

    segments = [
        {
            "key": "unchecked",
            "value": counts["unchecked"],
            "label": "未確認",
            "meaning": "未確認の割合",
            "color": "#D8CFC1",
        },
        {
            "key": "needsfix",
            "value": counts["needsfix"],
            "label": "要修正",
            "meaning": "要修正の割合",
            "color": "#D67A59",
        },
        {
            "key": "confirmed",
            "value": counts["confirmed"],
            "label": "確認済み",
            "meaning": "確認済みの割合",
            "color": "#6078A8",
        },
        {
            "key": "revised",
            "value": counts["revised"],
            "label": "修正済み",
            "meaning": "修正済みの割合",
            "color": "#2E8B7D",
        },
    ]

    # 実際の割合が小さい状態でも、色の層が消えないように表示用の高さだけ最小値を持たせる。
    # tooltip の割合・実数は実データをそのまま出す。
    positive_segments = [item for item in segments if int(item["value"]) > 0]
    raw_heights = []
    for item in positive_segments:
        raw_px = bar_height_px * pct(int(item["value"])) / 100
        raw_heights.append(max(raw_px, float(min_visible_px)))

    overflow = sum(raw_heights) - bar_height_px
    if overflow > 0:
        reducible = [max(0.0, h - min_visible_px) for h in raw_heights]
        total_reducible = sum(reducible)
        if total_reducible > 0:
            raw_heights = [
                h - overflow * (r / total_reducible)
                for h, r in zip(raw_heights, reducible)
            ]
        else:
            raw_heights = [bar_height_px / len(raw_heights)] * len(raw_heights)

    # 小数丸めで最後に隙間が出ないよう、最後の層だけ残り高さにする。
    visual_heights = []
    used = 0.0
    for idx, h in enumerate(raw_heights):
        if idx == len(raw_heights) - 1:
            final_h = max(1.0, bar_height_px - used)
        else:
            final_h = max(1.0, h)
            used += final_h
        visual_heights.append(final_h)

    segment_html = []
    if total <= 0:
        segment_html.append(f'<div class="seg empty" style="height:{bar_height_px}px" title="データなし"></div>')
    else:
        for item, visual_px in zip(positive_segments, visual_heights):
            value = int(item["value"])
            percent = pct(value)
            tooltip = f'{item["label"]}\n{item["meaning"]}\n全体の {percent:.1f}% ／ {value}件'
            segment_html.append(
                '<div '
                f'class="seg {item["key"]}" '
                f'style="height:{visual_px:.3f}px; background:{item["color"]};" '
                f'title="{html.escape(tooltip, quote=True)}" '
                f'data-tip="{html.escape(tooltip, quote=True)}">'
                '</div>'
            )

    progress_html = f"""
<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>
  html, body {{
    margin: 0;
    padding: 0;
    background: transparent;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    color: #26272F;
  }}
  .card {{
    box-sizing: border-box;
    width: 100%;
    min-height: 390px;
    padding: 14px 14px 18px;
    border: 1px solid #e3ded4;
    border-radius: 18px;
    background: rgba(255,253,247,.94);
    box-shadow: 0 12px 26px rgba(45,36,36,.07);
    overflow: visible;
  }}
  .head {{
    display: flex;
    align-items: baseline;
    justify-content: space-between;
    gap: 8px;
    margin-bottom: 14px;
  }}
  .title {{
    font-size: 15px;
    font-weight: 900;
    letter-spacing: .02em;
    color: #25243D;
  }}
  .percent {{
    font-size: 18px;
    font-weight: 900;
    color: #2E8B7D;
    font-variant-numeric: tabular-nums;
  }}
  .bar-wrap {{
    display: flex;
    justify-content: center;
    align-items: center;
    overflow: visible;
  }}
  .bar {{
    position: relative;
    width: 96px;
    height: {bar_height_px}px;
    display: flex;
    flex-direction: column;
    border-radius: 20px;
    overflow: visible;
    background: #EDE5D8;
    box-shadow:
      inset 0 0 0 1px rgba(255,255,255,.64),
      0 10px 22px rgba(45,36,36,.13);
  }}
  .bar::before {{
    content: "";
    position: absolute;
    inset: 0;
    border-radius: 20px;
    box-shadow: inset 0 0 0 1px rgba(24,33,43,.08);
    pointer-events: none;
    z-index: 3;
  }}
  .seg {{
    position: relative;
    width: 100%;
    flex: 0 0 auto;
    cursor: help;
    min-height: 1px;
  }}
  .seg:first-child {{ border-radius: 20px 20px 0 0; }}
  .seg:last-child {{ border-radius: 0 0 20px 20px; }}
  .seg:only-child {{ border-radius: 20px; }}
  .seg.empty {{
    height: 100%;
    background: #EDE5D8;
    border-radius: 20px;
  }}
  .seg:hover {{
    filter: brightness(.96) saturate(1.08);
    z-index: 5;
  }}
  .seg:hover::after {{
    content: attr(data-tip);
    position: fixed;
    left: 50%;
    top: 52%;
    transform: translate(-50%, -50%);
    box-sizing: border-box;
    width: 188px;
    padding: 9px 11px;
    border-radius: 12px;
    background: rgba(37, 36, 61, .96);
    color: #fffdf7;
    font-size: 12px;
    line-height: 1.45;
    white-space: pre-line;
    word-break: normal;
    text-align: left;
    box-shadow: 0 12px 26px rgba(31,41,51,.24);
    z-index: 9999;
    pointer-events: none;
  }}
</style>
</head>
<body>
  <div class="card">
    <div class="head">
      <span class="title">進捗</span>
      <span class="percent">{counts['percent']}%</span>
    </div>
    <div class="bar-wrap">
      <div class="bar" aria-label="進捗グラフ">
        {''.join(segment_html)}
      </div>
    </div>
  </div>
</body>
</html>
"""
    components.html(progress_html, height=404, scrolling=False)

def overview_dataframe(rows):
    return pd.DataFrame(
        [
            {
                "ID": row["ID"],
                "章": row["章"],
                "章内ID": row["章内ID"],
                "状態": row["状態"],
                "用語": "あり" if has_glossary_hit(row) else "",
                "英文プレビュー": truncate_text(row.get("英文", ""), 105),
                "訳文プレビュー": truncate_text(get_final_translation(row), 105),
                "メモ": truncate_text(row.get("メモ", ""), 70),
            }
            for row in rows
        ]
    )


def filter_rows(rows, selected_sections, selected_statuses, keyword, only_glossary):
    keyword = str(keyword or "").lower().strip()
    filtered = []
    for row in rows:
        if selected_sections and row.get("章") not in selected_sections:
            continue
        if selected_statuses and row.get("状態") not in selected_statuses:
            continue
        if only_glossary and not has_glossary_hit(row):
            continue
        if keyword:
            haystack = "\n".join(
                [
                    str(row.get("英文", "")),
                    str(row.get("原訳", "")),
                    str(row.get("修正訳", "")),
                    str(row.get("メモ", "")),
                ]
            ).lower()
            if keyword not in haystack:
                continue
        filtered.append(row)
    return filtered


# ============================================================
# ハイライトHTML
# ============================================================
def make_glossary_tooltip(item):
    """用語辞典ハイライト用の小さなHTMLツールチップを作る。

    以前は title / data-tip 属性に改行や <br> 相当の文字列を入れていたため、
    環境によっては「<br>」がそのまま見えてしまうことがあった。
    ここではツールチップ本体をHTML要素として持たせ、行ごとはCSSで縦並びにする。
    """
    rows = []
    english = clean_phrase(item.get("english", ""))
    japanese = clean_phrase(item.get("japanese", ""))
    category = clean_phrase(item.get("category", ""))
    note = str(item.get("note", "")).strip()

    if english:
        rows.append(("英語", english))
    if japanese:
        rows.append(("日本語訳", japanese))
    if category:
        rows.append(("カテゴリ", category))
    if note:
        rows.append(("メモ", note))

    if not rows:
        rows.append(("用語辞典", "登録済み"))

    line_html = []
    for label, value in rows:
        line_html.append(
            '<span class="tooltip-line">'
            f'<span class="tooltip-label">{html.escape(label)}：</span>'
            f'<span class="tooltip-value">{html.escape(value)}</span>'
            '</span>'
        )
    return '<span class="glossary-tooltip" role="tooltip">' + ''.join(line_html) + '</span>'


def highlight_english_terms(text, glossary):
    text = str(text or "")
    terms = []
    for item in glossary:
        term = clean_phrase(item.get("english", ""))
        if not term:
            continue
        terms.append((term, item))

    if not terms:
        return html.escape(text).replace("\n", "<br>")

    terms.sort(key=lambda pair: len(pair[0]), reverse=True)
    pattern = "|".join(f"({make_term_pattern(term)})" for term, _ in terms)

    item_by_lower = {}
    for term, item in terms:
        item_by_lower.setdefault(term.lower(), item)

    result = []
    last = 0
    for match in re.finditer(pattern, text, flags=re.IGNORECASE):
        start, end = match.span()
        if start < last:
            continue
        matched_text = text[start:end]
        result.append(html.escape(text[last:start]))

        key = matched_text.lower()
        item = item_by_lower.get(key)
        if item is None:
            for term, candidate in terms:
                if matched_text.lower() == term.lower():
                    item = candidate
                    break
        tooltip_html = make_glossary_tooltip(item or {})
        result.append(
            f'<span class="hl-glossary">{html.escape(matched_text)}{tooltip_html}</span>'
        )
        last = end

    result.append(html.escape(text[last:]))
    return "".join(result).replace("\n", "<br>")


def build_highlight_html(row, glossary, height=DEFAULT_VIEWER_HEIGHT):
    en_html = highlight_english_terms(row.get("英文", ""), glossary)
    ja_html = html.escape(get_final_translation(row)).replace("\n", "<br>")
    glossary_hits = sum(1 for item in glossary if clean_phrase(item.get("english", "")) and re.search(make_term_pattern(item.get("english", "")), str(row.get("英文", "")), flags=re.IGNORECASE))

    return f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{
    margin:0;
    font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;
    color:#172033;
    background:#ffffff;
}}
.wrap {{ padding: 0; }}
.legend {{
    display:flex;
    flex-wrap:wrap;
    gap:6px;
    margin-bottom:10px;
}}
.badge {{
    display:inline-block;
    padding:5px 10px;
    border-radius:999px;
    border:1px solid #dbe3ef;
    background:#f8fafc;
    font-size:12px;
    font-weight:700;
}}
.viewer {{
    display:grid;
    grid-template-columns:1fr 1fr;
    gap:12px;
    height:{height}px;
}}
.pane {{
    border:1px solid #dbe3ef;
    border-radius:18px;
    padding:16px 18px;
    overflow-y:auto;
    line-height:1.85;
    font-size:15.5px;
    background:linear-gradient(180deg,#ffffff 0%,#fbfdff 100%);
}}
.pane h3 {{
    position:sticky;
    top:-16px;
    z-index:3;
    background:#ffffff;
    padding:12px 0 10px 0;
    margin:-16px 0 12px 0;
    border-bottom:1px solid #eef2f7;
    color:#05285f;
}}
.hl-glossary {{
    position: relative;
    background:#dcfce7;
    border-bottom:2px solid #16a34a;
    border-radius:5px;
    padding:1px 4px;
    cursor:help;
    font-weight: 700;
}}
.glossary-tooltip {{
    display:none;
    position:absolute;
    left:0;
    top:1.9em;
    min-width:240px;
    max-width:360px;
    white-space:normal;
    background:#0f172a;
    color:#ffffff;
    padding:10px 12px;
    border-radius:12px;
    box-shadow:0 10px 24px rgba(15,23,42,.22);
    font-size:12.5px;
    line-height:1.55;
    font-weight:500;
    z-index:1000;
    pointer-events:none;
}}
.hl-glossary:hover .glossary-tooltip {{
    display:block;
}}
.tooltip-line {{
    display:block;
    margin:2px 0;
}}
.tooltip-label {{
    font-weight:800;
    color:#bbf7d0;
}}
.tooltip-value {{
    overflow-wrap:anywhere;
}}
</style>
</head>
<body>
<div class="wrap">
  <div class="legend">
    <span class="badge">緑：用語辞典ハイライト</span>
    <span class="badge">このペアのヒット数：{glossary_hits}</span>
    <span class="badge">カーソルを合わせると訳・カテゴリ・メモを表示</span>
  </div>
  <div class="viewer">
    <div class="pane"><h3>Original English</h3>{en_html}</div>
    <div class="pane"><h3>Japanese Translation</h3>{ja_html}</div>
  </div>
</div>
</body>
</html>
"""


def build_english_highlight_html(row, glossary, height=READ_EDITOR_HEIGHT):
    en_html = highlight_english_terms(row.get("英文", ""), glossary)
    glossary_hits = sum(
        1
        for item in glossary
        if clean_phrase(item.get("english", ""))
        and re.search(make_term_pattern(item.get("english", "")), str(row.get("英文", "")), flags=re.IGNORECASE)
    )
    return f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{
    margin:0;
    font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;
    color:#26272F;
    background:transparent;
}}
.pane {{
    box-sizing:border-box;
    height:{height}px;
    border:1px solid #E8DCCB;
    border-radius:18px;
    padding:16px 18px;
    overflow-y:auto;
    line-height:1.86;
    font-size:15.5px;
    background:linear-gradient(180deg,#FFFDF7 0%,#FBF7EF 100%);
    box-shadow:0 10px 24px rgba(45,36,36,.05);
}}
.pane h3 {{
    position:sticky;
    top:-16px;
    z-index:3;
    background:#FFFDF7;
    padding:12px 0 10px 0;
    margin:-16px 0 12px 0;
    border-bottom:1px solid #EFE3D1;
    color:#25243D;
    font-size:16px;
}}
.hit-count {{
    display:inline-block;
    margin-left:8px;
    padding:3px 8px;
    border-radius:999px;
    background:#EEF7F3;
    border:1px solid #C8E1D7;
    color:#2E8B7D;
    font-size:11px;
    font-weight:800;
}}
.hl-glossary {{
    position: relative;
    background:#E7F1E7;
    border-bottom:2px solid #6E9F84;
    border-radius:5px;
    padding:1px 4px;
    cursor:help;
    font-weight: 700;
}}
.glossary-tooltip {{
    display:none;
    position:absolute;
    left:0;
    top:1.9em;
    min-width:240px;
    max-width:360px;
    white-space:normal;
    background:#25243D;
    color:#FFFDF7;
    padding:10px 12px;
    border-radius:12px;
    box-shadow:0 10px 24px rgba(45,36,36,.22);
    font-size:12.5px;
    line-height:1.55;
    font-weight:500;
    z-index:1000;
    pointer-events:none;
}}
.hl-glossary:hover .glossary-tooltip {{ display:block; }}
.tooltip-line {{ display:block; margin:2px 0; }}
.tooltip-label {{ font-weight:850; color:#CDE5D7; }}
.tooltip-value {{ overflow-wrap:anywhere; }}
</style>
</head>
<body>
  <div id="en-pane" class="pane">
    <h3>Original English <span class="hit-count">用語 {glossary_hits}</span></h3>
    {en_html}
  </div>
<script>
(function() {{
  const pane = document.getElementById("en-pane");
  let attached = false;
  let lock = false;

  function findJapaneseTextarea() {{
    try {{
      const doc = window.parent.document;
      const areas = Array.from(doc.querySelectorAll("textarea"));
      return areas.find(el => (el.getAttribute("aria-label") || "").trim() === "Japanese Translation");
    }} catch (e) {{
      return null;
    }}
  }}

  function ratioOf(el) {{
    const max = Math.max(1, el.scrollHeight - el.clientHeight);
    return Math.max(0, Math.min(1, el.scrollTop / max));
  }}

  function setRatio(el, ratio) {{
    const max = Math.max(0, el.scrollHeight - el.clientHeight);
    el.scrollTop = max * ratio;
  }}

  function attachSync() {{
    if (attached) return true;
    const ja = findJapaneseTextarea();
    if (!ja) return false;
    attached = true;

    pane.addEventListener("scroll", function() {{
      if (lock) return;
      lock = true;
      setRatio(ja, ratioOf(pane));
      window.setTimeout(() => {{ lock = false; }}, 24);
    }}, {{ passive: true }});

    ja.addEventListener("scroll", function() {{
      if (lock) return;
      lock = true;
      setRatio(pane, ratioOf(ja));
      window.setTimeout(() => {{ lock = false; }}, 24);
    }}, {{ passive: true }});
    return true;
  }}

  const timer = window.setInterval(function() {{
    if (attachSync()) window.clearInterval(timer);
  }}, 250);
  attachSync();
}})();
</script>
</body>
</html>
"""


def inject_read_edit_helpers():
    """読む・修正画面のボタン色を、表示テキストに合わせて軽く調整する。"""
    helper_html = """
<script>
(function(){
  function styleButtons(){
    try {
      const doc = window.parent.document;
      doc.querySelectorAll('button').forEach(function(btn){
        const text = (btn.innerText || '').trim();
        if (text.includes('確認済みにして次へ')) {
          btn.style.background = '#2E8B7D';
          btn.style.borderColor = '#2E8B7D';
          btn.style.color = '#FFFDF7';
          btn.style.boxShadow = '0 10px 22px rgba(46,139,125,.18)';
        }
        if (text.includes('要修正にして次へ')) {
          btn.style.background = '#B07368';
          btn.style.borderColor = '#B07368';
          btn.style.color = '#FFFDF7';
          btn.style.boxShadow = '0 10px 22px rgba(176,115,104,.18)';
        }
      });
    } catch(e) {}
  }
  styleButtons();
  window.setTimeout(styleButtons, 300);
  window.setTimeout(styleButtons, 900);
})();
</script>
"""
    components.html(helper_html, height=1, scrolling=False)


# ============================================================
# 前後文脈・編集操作
# ============================================================
def get_context_rows(rows, row_id, radius):
    index = get_row_index_by_id(rows, row_id)
    if index is None:
        return []
    return rows[max(0, index - radius): min(len(rows), index + radius + 1)]


def render_context(rows, selected_id, radius):
    context = get_context_rows(rows, selected_id, radius)
    for row in context:
        is_current = int(row["ID"]) == int(selected_id)
        bg = "#eef4ff" if is_current else "#ffffff"
        border = "#003f8e" if is_current else "#dbe3ef"
        label = "現在" if is_current else "前後"
        st.markdown(
            f"""
<div style="border:1px solid {border}; background:{bg}; border-radius:14px; padding:12px 14px; margin:8px 0;">
  <div style="font-size:13px; color:#475569; font-weight:800; margin-bottom:6px;">
    {label} ／ ID {row['ID']} ／ {html.escape(str(row.get('章','')))}-{row.get('章内ID')} ／ {html.escape(str(row.get('状態','')))}
  </div>
  <div style="display:grid; grid-template-columns:1fr 1fr; gap:14px;">
    <div style="line-height:1.7;"><b>英文</b><br>{html.escape(str(row.get('英文','')))}</div>
    <div style="line-height:1.7;"><b>訳文</b><br>{html.escape(get_final_translation(row))}</div>
  </div>
</div>
""",
            unsafe_allow_html=True,
        )


def next_sequential_id(rows, current_id):
    index = get_row_index_by_id(rows, current_id)
    if index is None:
        return None
    if index + 1 < len(rows):
        return int(rows[index + 1]["ID"])
    return int(rows[index]["ID"])


def previous_sequential_id(rows, current_id):
    index = get_row_index_by_id(rows, current_id)
    if index is None:
        return None
    if index - 1 >= 0:
        return int(rows[index - 1]["ID"])
    return int(rows[index]["ID"])


def next_unchecked_id(rows, current_id):
    if not rows:
        return None
    index = get_row_index_by_id(rows, current_id)
    if index is None:
        index = -1
    order = list(range(index + 1, len(rows))) + list(range(0, index + 1))
    for i in order:
        if normalize_status(rows[i].get("状態")) in ["未確認", "要修正"]:
            return int(rows[i]["ID"])
    return next_sequential_id(rows, current_id)


def save_row(row_id, edited_translation, status, memo, restore=False):
    rows = st.session_state.get("rows", [])
    index = get_row_index_by_id(rows, row_id)
    if index is None:
        return False
    rows[index]["修正訳"] = str(rows[index].get("原訳", "")) if restore else str(edited_translation)
    rows[index]["状態"] = normalize_status(status)
    rows[index]["メモ"] = str(memo or "")
    rows[index]["更新日時"] = now_string()
    st.session_state["rows"] = rows
    return True


def insert_empty_row_after(rows, row_id):
    index = get_row_index_by_id(rows, row_id)
    if index is None:
        return rows
    base = rows[index]
    rows.insert(
        index + 1,
        {
            "ID": 0,
            "章": base.get("章", "Section"),
            "章内ID": 0,
            "英文": "",
            "原訳": "",
            "修正訳": "",
            "状態": "未確認",
            "メモ": "",
            "更新日時": now_string(),
        },
    )
    return renumber_rows(rows)


def delete_row_by_id(rows, row_id):
    index = get_row_index_by_id(rows, row_id)
    if index is None:
        return rows
    del rows[index]
    return renumber_rows(rows)


# ============================================================
# CSV / Word出力
# ============================================================
def make_csv_data(rows):
    df = pd.DataFrame(
        [
            {
                "ID": row["ID"],
                "章": row["章"],
                "章内ID": row["章内ID"],
                "状態": row["状態"],
                "英文": row.get("英文", ""),
                "原訳": row.get("原訳", ""),
                "修正訳": get_final_translation(row),
                "メモ": row.get("メモ", ""),
                "更新日時": row.get("更新日時", ""),
            }
            for row in rows
        ]
    )
    return df.to_csv(index=False).encode("utf-8-sig")


def make_docx_data(project_title, rows, export_style):
    document = Document()
    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"].font.size = Pt(10.5)

    document.add_heading(project_title or APP_NAME, level=1)
    document.add_paragraph(f"Exported from {APP_NAME} / {APP_VERSION}")
    document.add_paragraph(f"Saved at: {now_string()}")

    if export_style == "2列表（元文｜訳文）":
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "元文"
        table.rows[0].cells[1].text = "訳文"
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row.get("英文", ""))
            cells[1].text = get_final_translation(row)

    elif export_style == "縦並び（元文→訳文）":
        current_section = None
        for row in rows:
            if row.get("章") != current_section:
                current_section = row.get("章")
                document.add_heading(str(current_section), level=2)
            document.add_heading(f"ID {row.get('ID')}", level=3)
            document.add_paragraph("元文").runs[0].bold = True
            document.add_paragraph(str(row.get("英文", "")))
            document.add_paragraph("訳文").runs[0].bold = True
            document.add_paragraph(get_final_translation(row))

    elif export_style == "訳文のみ":
        current_section = None
        for row in rows:
            if row.get("章") != current_section:
                current_section = row.get("章")
                document.add_heading(str(current_section), level=2)
            document.add_paragraph(get_final_translation(row))

    else:
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["ID", "章", "状態", "元文", "訳文", "メモ"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for row in rows:
            values = [
                row["ID"],
                row["章"],
                row["状態"],
                row.get("英文", ""),
                get_final_translation(row),
                row.get("メモ", ""),
            ]
            cells = table.add_row().cells
            for i, value in enumerate(values):
                cells[i].text = str(value)

        if st.session_state.get("glossary"):
            document.add_page_break()
            document.add_heading("用語辞典", level=2)
            glossary_table = document.add_table(rows=1, cols=4)
            glossary_table.style = "Table Grid"
            for i, header in enumerate(["英語", "日本語", "カテゴリ", "メモ"]):
                glossary_table.rows[0].cells[i].text = header
            for item in st.session_state.get("glossary", []):
                cells = glossary_table.add_row().cells
                cells[0].text = str(item.get("english", ""))
                cells[1].text = str(item.get("japanese", ""))
                cells[2].text = str(item.get("category", ""))
                cells[3].text = str(item.get("note", ""))

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream.getvalue()


# ============================================================
# 画面：サイドバー
# ============================================================
def render_sidebar():
    rows = st.session_state.get("rows", [])
    counts = progress_counts(rows)

    with st.sidebar:
        st.header("作業パネル")

        render_sidebar_progress_card(rows)
        render_focus_timer()

        st.download_button(
            "JSONを保存",
            data=make_json_data(st.session_state["project_title"], st.session_state["unit"], rows),
            file_name=f"{safe_filename(st.session_state['project_title'])}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json",
            disabled=(not rows and not st.session_state.get("glossary")),
            use_container_width=True,
        )

        st.markdown("---")
        page_options = ["一覧・選択", "読む・修正", "用語辞典", "出力"]
        if st.session_state.get("active_page") not in page_options:
            st.session_state["active_page"] = "一覧・選択"
        st.session_state["active_page"] = st.radio(
            "画面",
            page_options,
            index=page_options.index(st.session_state["active_page"]),
        )

        st.markdown("---")
        c1, c2 = st.columns(2)
        with c1:
            if st.button("設定", use_container_width=True):
                st.session_state["show_settings"] = not st.session_state["show_settings"]
        with c2:
            if st.button("入力", use_container_width=True):
                st.session_state["show_json_loader"] = True
                st.session_state["show_text_adder"] = True

        if st.session_state.get("show_settings", False):
            st.markdown("### 設定")
            st.session_state["project_title"] = st.text_input(
                "プロジェクト名",
                value=st.session_state.get("project_title", "paper_parallel_project"),
            )
            st.session_state["unit"] = st.radio(
                "処理単位",
                ["段落", "文"],
                index=0 if st.session_state.get("unit", "段落") == "段落" else 1,
                help="10000 words級の論文では、まず段落単位がおすすめです。",
            )
            st.session_state["page_size"] = st.selectbox(
                "1ページの表示件数",
                [10, 20, 50, 100, 200],
                index=[10, 20, 50, 100, 200].index(st.session_state.get("page_size", 20)),
            )
            st.session_state["context_radius"] = st.slider(
                "前後文脈の表示範囲",
                min_value=1,
                max_value=5,
                value=int(st.session_state.get("context_radius", 2)),
            )


# ============================================================
# 画面：読み込み・テキスト追加
# ============================================================
def render_input_area():
    show_any = st.session_state.get("show_json_loader", False) or st.session_state.get("show_text_adder", False)
    if not show_any:
        c1, c2 = st.columns(2)
        with c1:
            if st.button("JSON読み込み欄を開く", use_container_width=True):
                st.session_state["show_json_loader"] = True
                st.rerun()
        with c2:
            if st.button("テキスト追加欄を開く", use_container_width=True):
                st.session_state["show_text_adder"] = True
                st.rerun()
        return

    setup_left, setup_right = st.columns(2)

    with setup_left:
        if st.session_state.get("show_json_loader", False):
            with st.expander("保存済みJSONから作業を再開", expanded=True):
                uploaded_json = st.file_uploader(
                    "JSONファイルを選択してください",
                    type=["json"],
                    help="この版のJSONだけでなく、旧 alignment_results 形式もできる範囲で読み込めます。",
                )
                if uploaded_json is not None:
                    st.write(f"選択中：{uploaded_json.name}")
                    if st.button("このJSONを読み込む", type="primary", use_container_width=True):
                        try:
                            project_title, unit, rows = load_rows_from_json(uploaded_json)
                            st.session_state["project_title"] = project_title
                            st.session_state["unit"] = unit if unit in ["段落", "文"] else "段落"
                            st.session_state["selected_id"] = rows[0]["ID"] if rows else None
                            st.session_state["active_page"] = "一覧・選択"
                            st.session_state["show_json_loader"] = False
                            st.session_state["show_text_adder"] = False
                            st.session_state["show_settings"] = False
                            st.success("JSONを読み込みました。入力欄と設定欄は閉じました。")
                            st.rerun()
                        except Exception as exc:
                            st.error(f"読み込みに失敗しました：{exc}")
                if st.button("JSON読み込み欄を閉じる", use_container_width=True):
                    st.session_state["show_json_loader"] = False
                    st.rerun()

    with setup_right:
        if st.session_state.get("show_text_adder", False):
            with st.expander("英文・訳文を追加", expanded=True):
                st.caption(f"現在の分割単位：{st.session_state.get('unit', '段落')}。変更したい場合は左の設定を開いてください。")
                section_title = st.text_input("章・セクション名", value="Section")
                english_text = st.text_area("英文", height=220)
                japanese_text = st.text_area("訳文", height=220)
                add_mode = st.radio(
                    "追加方法",
                    ["現在の作業に追加", "新規作成（現在の作業を置き換え）"],
                    index=0,
                )
                if st.button("テキストを反映", type="primary", use_container_width=True):
                    en_segments = segment_text(english_text, "英語", st.session_state["unit"])
                    ja_segments = segment_text(japanese_text, "日本語", st.session_state["unit"])
                    if not en_segments and not ja_segments:
                        st.warning("英文または訳文を入力してください。")
                    else:
                        existing_rows = [] if add_mode.startswith("新規作成") else st.session_state.get("rows", [])
                        new_rows = make_rows_from_segments(
                            en_segments,
                            ja_segments,
                            section_title,
                            start_id=len(existing_rows) + 1,
                        )
                        st.session_state["rows"] = renumber_rows(existing_rows + new_rows)
                        if st.session_state["rows"] and st.session_state.get("selected_id") is None:
                            st.session_state["selected_id"] = st.session_state["rows"][0]["ID"]
                        st.session_state["active_page"] = "一覧・選択"
                        st.session_state["show_json_loader"] = False
                        st.session_state["show_text_adder"] = False
                        st.session_state["show_settings"] = False
                        st.success("テキストを追加しました。入力欄と設定欄は閉じました。")
                        st.rerun()
                if st.button("テキスト追加欄を閉じる", use_container_width=True):
                    st.session_state["show_text_adder"] = False
                    st.rerun()


# ============================================================
# 画面：一覧・選択
# ============================================================
def render_overview_page():
    rows = st.session_state.get("rows", [])
    if not rows:
        st.info("まずJSONを読み込むか、英文・訳文を追加してください。")
        return

    st.markdown("### 一覧・選択")
    filter_left, filter_mid, filter_right = st.columns([1, 1, 1])
    sections = sorted({str(row.get("章", "")) for row in rows})
    with filter_left:
        selected_sections = st.multiselect("章で絞り込み", sections)
    with filter_mid:
        selected_statuses = st.multiselect("状態で絞り込み", STATUS_OPTIONS)
    with filter_right:
        keyword = st.text_input("キーワード検索")
    only_glossary = st.checkbox("用語辞典に登録した英語が出てくるペアだけ表示")

    filtered = filter_rows(rows, selected_sections, selected_statuses, keyword, only_glossary)
    st.session_state["last_filtered_ids"] = [int(row["ID"]) for row in filtered]

    st.markdown(f"<span class='ppr-chip ppr-chip-blue'>表示対象：{len(filtered)} / {len(rows)}</span>", unsafe_allow_html=True)

    if not filtered:
        st.warning("条件に合うペアがありません。")
        return

    page_size = int(st.session_state.get("page_size", 20))
    total_pages = max(1, math.ceil(len(filtered) / page_size))
    page_number = st.number_input("ページ", min_value=1, max_value=total_pages, value=1, step=1)
    start = (int(page_number) - 1) * page_size
    end = start + page_size
    page_rows = filtered[start:end]

    overview_df = overview_dataframe(page_rows)
    table_event = st.dataframe(
        overview_df,
        use_container_width=True,
        hide_index=True,
        on_select="rerun",
        selection_mode="single-row",
        key=f"overview_table_{int(page_number)}_{page_size}_{len(filtered)}",
    )
    selected_table_rows = []
    if hasattr(table_event, "selection") and hasattr(table_event.selection, "rows"):
        selected_table_rows = table_event.selection.rows
    elif isinstance(table_event, dict):
        selected_table_rows = table_event.get("selection", {}).get("rows", [])

    if selected_table_rows:
        picked_index = int(selected_table_rows[0])
        if 0 <= picked_index < len(overview_df):
            st.session_state["selected_id"] = int(overview_df.iloc[picked_index]["ID"])

    choices = [int(row["ID"]) for row in page_rows]
    current_id = st.session_state.get("selected_id")
    if current_id not in choices:
        current_id = choices[0]
        st.session_state["selected_id"] = current_id

    # 表をクリックした選択を、この selectbox にも確実に反映させる。
    selector_key = "read_pair_selector"
    if st.session_state.get(selector_key) not in choices or st.session_state.get(selector_key) != current_id:
        st.session_state[selector_key] = current_id

    selected_id = st.selectbox(
        "読むペアを選ぶ",
        choices,
        key=selector_key,
        format_func=lambda rid: f"ID {rid}: {truncate_text(get_row_by_id(rows, rid).get('英文', ''), 80)}",
    )
    st.session_state["selected_id"] = int(selected_id)

    if st.button("選択して読む", type="primary"):
        st.session_state["selected_id"] = int(selected_id)
        st.session_state["active_page"] = "読む・修正"
        st.rerun()


# ============================================================
# 画面：読む・修正
# ============================================================
def render_read_edit_page():
    rows = st.session_state.get("rows", [])
    if not rows:
        st.info("まずJSONを読み込むか、英文・訳文を追加してください。")
        return

    if st.session_state.get("selected_id") is None:
        st.session_state["selected_id"] = rows[0]["ID"]

    selected_row = get_row_by_id(rows, st.session_state["selected_id"])
    if selected_row is None:
        st.session_state["selected_id"] = rows[0]["ID"]
        selected_row = rows[0]

    st.markdown("### 読む・修正")
    current_index = get_row_index_by_id(rows, selected_row["ID"])

    top_left, top_right = st.columns([1.2, 1])
    with top_left:
        glossary_hit = has_glossary_hit(selected_row)
        st.markdown(
            f"""
<div class="ppr-card">
  <b>選択中：</b>ID {selected_row['ID']} ／ {html.escape(str(selected_row.get('章','')))}-{selected_row.get('章内ID')}<br>
  <span class="ppr-chip ppr-chip-blue">状態：{html.escape(str(selected_row.get('状態','')))}</span>
  <span class="ppr-chip ppr-chip-green">用語辞典ヒット：{'あり' if glossary_hit else 'なし'}</span>
</div>
""",
            unsafe_allow_html=True,
        )
    with top_right:
        n1, n2, n3, n4 = st.columns(4)
        with n1:
            if st.button("前へ", disabled=(current_index is None or current_index <= 0), use_container_width=True):
                st.session_state["selected_id"] = previous_sequential_id(rows, selected_row["ID"])
                st.rerun()
        with n2:
            if st.button("次へ", disabled=(current_index is None or current_index >= len(rows) - 1), use_container_width=True):
                st.session_state["selected_id"] = next_sequential_id(rows, selected_row["ID"])
                st.rerun()
        with n3:
            if st.button("次の未確認", use_container_width=True):
                st.session_state["selected_id"] = next_unchecked_id(rows, selected_row["ID"])
                st.rerun()
        with n4:
            if st.button("一覧へ", use_container_width=True):
                st.session_state["active_page"] = "一覧・選択"
                st.rerun()

    translation_key = f"translation_editor_{selected_row['ID']}"
    memo_key = f"memo_editor_{selected_row['ID']}"
    if translation_key not in st.session_state:
        st.session_state[translation_key] = get_final_translation(selected_row)
    if memo_key not in st.session_state:
        st.session_state[memo_key] = str(selected_row.get("メモ", ""))

    left, right = st.columns([1, 1], gap="medium")
    with left:
        components.html(
            build_english_highlight_html(selected_row, st.session_state.get("glossary", []), height=READ_EDITOR_HEIGHT),
            height=READ_EDITOR_HEIGHT + 8,
            scrolling=False,
        )

    with right:
        st.text_area(
            "Japanese Translation",
            key=translation_key,
            height=READ_EDITOR_HEIGHT,
            help="ここに直接修正訳を書き込めます。英語欄とスクロール位置が同期します。",
        )

    with st.expander("前後文脈を見る", expanded=False):
        render_context(rows, selected_row["ID"], int(st.session_state.get("context_radius", 2)))

    st.text_area("メモ", key=memo_key, height=82)

    action_save, action_confirm, action_fix = st.columns([1, 1.2, 1.2])
    save_only = False
    confirm_and_next = False
    needsfix_and_next = False
    with action_save:
        save_only = st.button("このペアを保存", use_container_width=True)
    with action_confirm:
        confirm_and_next = st.button("確認済みにして次へ", type="primary", use_container_width=True)
    with action_fix:
        needsfix_and_next = st.button("要修正にして次へ", use_container_width=True)

    inject_read_edit_helpers()

    if save_only or confirm_and_next or needsfix_and_next:
        save_status = normalize_status(selected_row.get("状態"))
        if confirm_and_next:
            save_status = "確認済み"
        elif needsfix_and_next:
            save_status = "要修正"

        ok = save_row(
            selected_row["ID"],
            st.session_state.get(translation_key, ""),
            save_status,
            st.session_state.get(memo_key, ""),
            False,
        )
        if ok:
            if confirm_and_next or needsfix_and_next:
                st.session_state["selected_id"] = next_sequential_id(st.session_state["rows"], selected_row["ID"])
            st.success("保存しました。")
            st.rerun()
        else:
            st.error("保存できませんでした。")

    with st.expander("結合・削除・空ペア", expanded=False):
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("空ペア追加", use_container_width=True):
                st.session_state["rows"] = insert_empty_row_after(rows, selected_row["ID"])
                st.session_state["selected_id"] = next_sequential_id(st.session_state["rows"], selected_row["ID"])
                st.rerun()
        with c2:
            if st.button("次の訳文を結合", disabled=(current_index is None or current_index >= len(rows) - 1), use_container_width=True):
                current = rows[current_index]
                next_row = rows[current_index + 1]
                current["原訳"] = "\n\n".join(
                    [p for p in [str(current.get("原訳", "")).strip(), str(next_row.get("原訳", "")).strip()] if p]
                )
                current["修正訳"] = "\n\n".join(
                    [p for p in [get_final_translation(current), get_final_translation(next_row)] if p]
                )
                current["状態"] = "修正済み"
                current["更新日時"] = now_string()
                del rows[current_index + 1]
                st.session_state["rows"] = renumber_rows(rows)
                st.success("次の訳文を結合しました。")
                st.rerun()
        with c3:
            if st.button("このペアを削除", disabled=(len(rows) <= 1), use_container_width=True):
                next_id = next_sequential_id(rows, selected_row["ID"])
                st.session_state["rows"] = delete_row_by_id(rows, selected_row["ID"])
                if st.session_state["rows"]:
                    st.session_state["selected_id"] = next_id if get_row_by_id(st.session_state["rows"], next_id) else st.session_state["rows"][0]["ID"]
                else:
                    st.session_state["selected_id"] = None
                st.rerun()


# ============================================================
# 画面：用語辞典
# ============================================================
def render_glossary_page():
    st.markdown("### 用語辞典")
    st.caption("英語用語を登録すると、読む・修正画面の英文中で自動ハイライトされます。カーソルを合わせると訳・カテゴリ・メモが出ます。")

    with st.form("glossary_add_form", clear_on_submit=True):
        c1, c2 = st.columns(2)
        with c1:
            english = st.text_input("英語")
        with c2:
            japanese = st.text_input("日本語訳")
        c3, c4 = st.columns([1, 2])
        with c3:
            category = st.selectbox("カテゴリ", GLOSSARY_CATEGORIES)
        with c4:
            note = st.text_input("メモ")
        submitted = st.form_submit_button("用語を追加", type="primary")

    if submitted:
        ok, message = add_unique_glossary(english, japanese, category, note)
        if ok:
            st.success(message)
            st.rerun()
        else:
            st.warning(message)

    glossary = st.session_state.get("glossary", [])
    if not glossary:
        st.info("まだ用語が登録されていません。")
        return

    st.markdown("#### 登録済み用語")
    df = pd.DataFrame(glossary)
    display_columns = ["english", "japanese", "category", "note"]
    edited_df = st.data_editor(
        df[display_columns],
        use_container_width=True,
        num_rows="dynamic",
        hide_index=True,
        column_config={
            "english": "英語",
            "japanese": "日本語訳",
            "category": st.column_config.SelectboxColumn("カテゴリ", options=GLOSSARY_CATEGORIES),
            "note": "メモ",
        },
    )
    if st.button("用語辞典の編集を保存", type="primary"):
        new_glossary = []
        for _, row in edited_df.iterrows():
            english_text = clean_phrase(row.get("english", ""))
            japanese_text = clean_phrase(row.get("japanese", ""))
            if not english_text and not japanese_text:
                continue
            new_glossary.append(
                {
                    "id": len(new_glossary) + 1,
                    "english": english_text,
                    "japanese": japanese_text,
                    "category": str(row.get("category", "専門用語") or "専門用語"),
                    "note": str(row.get("note", "")),
                    "created_at": now_string(),
                }
            )
        st.session_state["glossary"] = new_glossary
        cleanup_glossary()
        st.success("用語辞典を保存しました。")
        st.rerun()


# ============================================================
# 画面：出力
# ============================================================
def render_export_page():
    rows = st.session_state.get("rows", [])
    st.markdown("### 出力")
    if not rows:
        st.info("出力するデータがまだありません。")
        return

    base_name = safe_filename(st.session_state.get("project_title", "paper_parallel"))
    c1, c2 = st.columns(2)
    with c1:
        st.download_button(
            "CSVを保存",
            data=make_csv_data(rows),
            file_name=f"{base_name}.csv",
            mime="text/csv",
            use_container_width=True,
        )
    with c2:
        export_style = st.selectbox("Word出力形式", WORD_EXPORT_OPTIONS)
        st.download_button(
            "Wordを保存",
            data=make_docx_data(st.session_state.get("project_title"), rows, export_style),
            file_name=f"{base_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    st.caption("JSON保存は左の作業パネルからいつでもできます。")



# ============================================================
# 集中タイマー
# ============================================================
def render_focus_timer():
    """作業パネル内に置く軽量タイマー。Streamlitの再実行を起こさないよう、JSだけで動かす。"""
    timer_html = """
<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>
  html, body {
    margin: 0;
    padding: 0;
    background: transparent;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    color: #25243D;
  }
  .timer-card {
    box-sizing: border-box;
    width: 100%;
    min-height: 126px;
    margin: 8px 0 10px;
    padding: 13px 13px 12px;
    border-radius: 18px;
    border: 1px solid rgba(232, 220, 203, .94);
    background:
      radial-gradient(circle at 84% 0%, rgba(200,190,221,.24) 0, rgba(200,190,221,0) 36%),
      rgba(255,253,247,.94);
    box-shadow: 0 10px 24px rgba(45,36,36,.07);
  }
  .timer-head {
    display: flex;
    align-items: baseline;
    justify-content: space-between;
    gap: 8px;
    margin-bottom: 8px;
  }
  .timer-label {
    font-size: 13px;
    font-weight: 900;
    letter-spacing: .02em;
    color: #25243D;
  }
  .timer-time {
    font-size: 25px;
    line-height: 1;
    font-weight: 920;
    letter-spacing: .02em;
    color: #2E8B7D;
    font-variant-numeric: tabular-nums;
  }
  .timer-actions {
    display: grid;
    grid-template-columns: 1fr 1fr 1fr;
    gap: 6px;
  }
  button {
    appearance: none;
    border: 1px solid #DDCBBB;
    border-radius: 12px;
    padding: 8px 0;
    font-size: 11px;
    font-weight: 850;
    color: #3A3447;
    background: #FFF9F1;
    cursor: pointer;
    transition: transform .12s ease, border-color .12s ease, background .12s ease;
  }
  button:hover {
    transform: translateY(-1px);
    border-color: #B07368;
    background: #FFF3EA;
  }
  button.primary {
    color: #FFF9EF;
    border-color: #2E8B7D;
    background: #2E8B7D;
  }
  button.primary:hover {
    background: #26786C;
    border-color: #26786C;
  }
</style>
</head>
<body>
  <div class="timer-card">
    <div class="timer-head">
      <div class="timer-label">集中タイマー</div>
      <div id="time" class="timer-time">00:00</div>
    </div>
    <div class="timer-actions">
      <button id="start" class="primary">Start</button>
      <button id="stop">Stop</button>
      <button id="reset">Reset</button>
    </div>
  </div>
<script>
(function(){
  const STORAGE_KEY = "ppr_focus_timer_v1";
  const timeEl = document.getElementById("time");
  const startBtn = document.getElementById("start");
  const stopBtn = document.getElementById("stop");
  const resetBtn = document.getElementById("reset");

  function loadState() {
    try {
      const saved = JSON.parse(localStorage.getItem(STORAGE_KEY) || "{}");
      return {
        running: Boolean(saved.running),
        startedAt: Number(saved.startedAt || Date.now()),
        elapsed: Number(saved.elapsed || 0)
      };
    } catch(e) {
      return { running: false, startedAt: Date.now(), elapsed: 0 };
    }
  }

  let state = loadState();

  function saveState() {
    localStorage.setItem(STORAGE_KEY, JSON.stringify(state));
  }

  function currentElapsed() {
    return state.running ? Math.max(0, Date.now() - state.startedAt) : Math.max(0, state.elapsed);
  }

  function render() {
    const totalSeconds = Math.floor(currentElapsed() / 1000);
    const minutes = Math.floor(totalSeconds / 60);
    const seconds = totalSeconds % 60;
    timeEl.textContent = String(minutes).padStart(2, "0") + ":" + String(seconds).padStart(2, "0");
    startBtn.textContent = state.running ? "Running" : "Start";
  }

  startBtn.addEventListener("click", function(){
    if (!state.running) {
      state.startedAt = Date.now() - state.elapsed;
      state.running = true;
      saveState();
      render();
    }
  });

  stopBtn.addEventListener("click", function(){
    if (state.running) {
      state.elapsed = currentElapsed();
      state.running = false;
      saveState();
      render();
    }
  });

  resetBtn.addEventListener("click", function(){
    state = { running: false, startedAt: Date.now(), elapsed: 0 };
    saveState();
    render();
  });

  render();
  setInterval(render, 250);
})();
</script>
</body>
</html>
"""
    components.html(timer_html, height=140, scrolling=False)

# ============================================================
# アプリ本体
# ============================================================
inject_style()
init_state()
cleanup_glossary()
st.markdown(
    f"""
<div class="ppr-hero">
  <h1>Paper Parallel Reader</h1>
  <p>用語辞典を中心に、論文の英文と和訳を軽く確認・修正するローカルエディタ</p>
  <span class="ppr-badge">Version: {APP_VERSION}</span>
</div>
""",
    unsafe_allow_html=True,
)

render_sidebar()
render_input_area()

active_page = st.session_state.get("active_page", "一覧・選択")
if active_page == "一覧・選択":
    render_overview_page()
elif active_page == "読む・修正":
    render_read_edit_page()
elif active_page == "用語辞典":
    render_glossary_page()
elif active_page == "出力":
    render_export_page()
